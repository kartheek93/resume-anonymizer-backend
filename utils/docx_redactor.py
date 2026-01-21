import re
from docx import Document
from utils.location_keywords import LOCATION_KEYWORDS

EMOJI_PATTERN = re.compile("[\U0001F300-\U0001FAFF\u2600-\u27BF]", re.UNICODE)
EMAIL_PATTERN = re.compile(r"@")
PHONE_PATTERN = re.compile(r"\+?\d[\d\s\-]{8,}\d")

DOB_PATTERN = re.compile(
    r"\b(dob|date\s+of\s+birth)\b|"
    r"\b\d{1,2}\s+(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\s+\d{4}\b",
    re.I
)

IDENTITY_LABEL_PATTERN = re.compile(
    r"\b(birth\s*place|birthplace|nationality|citizenship|gender|marital\s*status)\b",
    re.I
)

ADDRESS_PATTERN = re.compile(r"\d{6}|\d{1,4}[-/]\d{1,4}", re.I)

SOCIAL_LABEL_PATTERN = re.compile(
    r"\b(linked\s*in|github|git\s*hub|leet\s*code|hacker\s*rank|portfolio|profile)\b\s*[:\-|/]?\s*\S+",
    re.I
)

USERNAME_PATTERN = re.compile(r"\b[a-z][a-z0-9._-]{3,}\b", re.I)

SOCIAL_CONTEXT = [
    "linkedin", "github", "leetcode", "hackerrank", "portfolio", "profile"
]

def looks_like_name(text):
    t = text.strip()
    if not t:
        return False
    if EMAIL_PATTERN.search(t) or PHONE_PATTERN.search(t):
        return False
    if sum(c.isdigit() for c in t) > 2:
        return False
    words = t.replace("|", " ").split()
    return 1 <= len(words) <= 6


def remove_hyperlinks(paragraph):
    p = paragraph._p
    for link in p.findall(".//w:hyperlink", namespaces=p.nsmap):
        p.remove(link)


def remove_inline_images(paragraph):
    for run in paragraph.runs:
        if "graphicData" in run._element.xml:
            run.clear()


def redact_docx_inplace(input_path, output_path):
    doc = Document(input_path)

    total_paras = len(doc.paragraphs)
    header_limit = max(1, int(total_paras * 0.30))

    for idx, para in enumerate(doc.paragraphs):
        raw = para.text
        text = raw.lower()

        remove_hyperlinks(para)
        remove_inline_images(para)
        para.text = EMOJI_PATTERN.sub("", para.text)

        # ❌ Do nothing below top 30%
        if idx >= header_limit:
            continue

        delete_line = (
            EMAIL_PATTERN.search(raw)
            or PHONE_PATTERN.search(raw)
            or DOB_PATTERN.search(text)
            or IDENTITY_LABEL_PATTERN.search(text)
            or ADDRESS_PATTERN.search(text)
            or SOCIAL_LABEL_PATTERN.search(raw)
            or (any(k in text for k in SOCIAL_CONTEXT) and USERNAME_PATTERN.search(raw))
        )

        # Location only when tied to personal info
        if any(loc in text for loc in LOCATION_KEYWORDS) and (
            EMAIL_PATTERN.search(raw)
            or PHONE_PATTERN.search(raw)
            or SOCIAL_LABEL_PATTERN.search(raw)
            or ADDRESS_PATTERN.search(text)
        ):
            delete_line = True

        if delete_line:
            if looks_like_name(raw):
                words = raw.split()
                keep = words[:max(1, len(words)//2)]
                para.text = " ".join(keep)
            else:
                para.text = ""

    doc.save(output_path)
